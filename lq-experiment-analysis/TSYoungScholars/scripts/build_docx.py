#!/usr/bin/env python3
"""将工作计划书.md 排成贴近泰山青年模板的可编辑 Word 稿（A4 / 宋体 / 2.2·2.3 居中）。"""
from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Emu
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "工作计划书.md"
FIG = ROOT / "figures"
OUT = ROOT / "工作计划书.docx"

SONG = "Songti SC"
HEI = "Heiti SC"
EN = "Times New Roman"


def set_run_font(run, name_east=SONG, size_pt=12, bold=False, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = EN
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), EN)
    rFonts.set(qn("w:hAnsi"), EN)
    rFonts.set(qn("w:eastAsia"), name_east)
    rFonts.set(qn("w:cs"), EN)


def set_paragraph_format(p, *, first_line=False, center=False, space_after=6, space_before=0, line=22):
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line)
    pf.widow_control = True
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(0.74) if first_line else Cm(0)


def shade_run(run, hex_color="FFF2CC"):
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)


def add_bottom_border(paragraph, color="2F4F4F", sz="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_border(section, color="C5D5C5"):
    """浅色页框，呼应模板截图的淡绿底，但不把整页涂绿以免打印发灰。"""
    sectPr = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "16")
        el.set(qn("w:color"), color)
        pgBorders.append(el)
    sectPr.append(pgBorders)


def set_cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, **cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in cm.items():
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(int(val * 567)))  # cm → twips approx 567
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def add_runs_from_markdown(paragraph, text, *, size=12, font=SONG, bold=False):
    """支持 **bold**、*italic*、`code`、[待填…] 高亮。"""
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]\n]*待填[^\]\n]*\])"
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            set_run_font(r, font, size, bold=bold)
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            set_run_font(r, HEI if font == HEI else SONG, size, bold=True)
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            set_run_font(r, "Menlo", size, bold=False)
        elif token.startswith("[") and "待填" in token:
            r = paragraph.add_run(token)
            set_run_font(r, SONG, size, bold=True, color=RGBColor(0x8B, 0x1A, 0x1A))
            shade_run(r, "FFF2CC")
        else:
            r = paragraph.add_run(token[1:-1])
            set_run_font(r, SONG, size, italic=True)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_run_font(r, font, size, bold=bold)


def add_heading_text(doc, text, *, level, center=False):
    p = doc.add_paragraph()
    if level == 0:
        set_paragraph_format(p, center=True, space_after=10, space_before=0, line=28)
        r = p.add_run(text)
        set_run_font(r, HEI, 18, bold=True)
        add_bottom_border(p, "3D5C3D", "18")
    elif level == 1:
        set_paragraph_format(p, center=center, first_line=False, space_after=8, space_before=14, line=24)
        r = p.add_run(text)
        set_run_font(r, HEI, 14, bold=True, color=RGBColor(0x2F, 0x4F, 0x2F) if center else None)
        if center:
            add_bottom_border(p, "5B7A5B", "12")
    else:
        set_paragraph_format(p, center=False, first_line=False, space_after=6, space_before=10, line=22)
        r = p.add_run(text)
        set_run_font(r, HEI, 12, bold=True)
    return p


def add_body(doc, text, *, first_line=True, size=12):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line, space_after=4, line=22)
    add_runs_from_markdown(p, text, size=size)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, space_after=8, space_before=2, line=20)
    r = p.add_run(text)
    set_run_font(r, SONG, 10.5, italic=True, color=RGBColor(0x55, 0x55, 0x55))
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, space_after=10, space_before=4, line=20)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs_from_markdown(p, text, size=10.5)
    for run in p.runs:
        run.italic = False
    return p


def add_picture(doc, path: Path, width_cm=15.4):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, center=True, space_after=2, space_before=8, line=16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    return p


def parse_table(lines):
    rows = []
    for ln in lines:
        if re.match(r"^\|?\s*:?-{3,}", ln.replace("|", " | ")):
            # skip separator — handled below
            pass
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            continue
        rows.append([c.replace("<br>", "\n").replace("<br/>", "\n") for c in cells])
    return rows


def add_table(doc, rows):
    if not rows:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(2.4), Cm(8.6), Cm(5.0)] if ncols == 3 else [Cm(16.0 / ncols)] * ncols
    for row_i, row in enumerate(rows):
        prevent_row_split(table.rows[row_i])
        for col_i, val in enumerate(row):
            cell = table.cell(row_i, col_i)
            cell.width = widths[col_i]
            set_cell_margins(cell, top=0.08, bottom=0.08, left=0.12, right=0.12)
            if row_i == 0:
                set_cell_shading(cell, "D9E6D9")
            elif row_i % 2 == 0:
                set_cell_shading(cell, "F4F8F4")
            # clear default para
            cell.paragraphs[0].clear()
            parts = val.split("\n")
            for pi, part in enumerate(parts):
                para = cell.paragraphs[0] if pi == 0 else cell.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf = para.paragraph_format
                pf.space_after = Pt(2)
                pf.space_before = Pt(0)
                pf.line_spacing = Pt(18)
                pf.first_line_indent = Cm(0)
                add_runs_from_markdown(para, part, size=10.5, font=HEI if row_i == 0 else SONG)
                if row_i == 0:
                    for run in para.runs:
                        run.bold = True
    doc.add_paragraph()


def find_figure(name_fragment: str) -> Path | None:
    cands = sorted(FIG.glob(f"*{name_fragment}*"))
    pngs = [p for p in cands if p.suffix.lower() in {".png", ".jpg", ".jpeg"}]
    return pngs[0] if pngs else None


def convert():
    raw = MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    add_page_border(section)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(fp, first_line=False, center=True, space_after=0, line=16)
    r1 = fp.add_run("— ")
    set_run_font(r1, SONG, 9, color=RGBColor(0x66, 0x66, 0x66))
    r_page = fp.add_run()
    set_run_font(r_page, EN, 9, color=RGBColor(0x66, 0x66, 0x66))
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_page._r.append(fld_begin)
    r_page._r.append(instr)
    r_page._r.append(fld_end)
    r2 = fp.add_run(" —")
    set_run_font(r2, SONG, 9, color=RGBColor(0x66, 0x66, 0x66))

    # default style
    normal = doc.styles["Normal"]
    normal.font.name = EN
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), SONG)

    i = 0
    n = len(raw)
    while i < n:
        line = raw[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue
        if stripped.startswith("```"):
            i += 1
            while i < n and not raw[i].strip().startswith("```"):
                i += 1
            i += 1
            continue

        if stripped.startswith("# "):
            add_heading_text(doc, stripped[2:].strip(), level=0, center=True)
            i += 1
            continue

        if stripped.startswith("## "):
            title = stripped[3:].strip()
            center = title.startswith("2.2") or title.startswith("2.3")
            add_heading_text(doc, title, level=1, center=center)
            i += 1
            continue

        if stripped.startswith("### "):
            add_heading_text(doc, stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("> "):
            add_note(doc, stripped[2:].strip())
            i += 1
            continue

        if stripped.startswith("!["):
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
            if m:
                alt, rel = m.group(1), m.group(2)
                path = (ROOT / rel).resolve()
                if not path.exists():
                    # fallback by keyword
                    key = "mechanism" if "mechanism" in rel else "technical_route"
                    path = find_figure(key)
                if path and path.exists():
                    add_picture(doc, path)
                else:
                    add_note(doc, f"[图未找到：{rel}]")
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < n and re.search(r"---", raw[i + 1]):
            block = [stripped]
            i += 1
            while i < n and raw[i].strip().startswith("|"):
                block.append(raw[i].strip())
                i += 1
            add_table(doc, parse_table(block))
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False, space_after=3, line=22)
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.37)
            add_runs_from_markdown(p, "·  " + stripped[2:].strip(), size=12)
            i += 1
            continue

        # caption-like figure sentences
        if stripped.startswith("**图"):
            add_caption(doc, stripped)
            i += 1
            continue

        # parenthetical template notes
        if stripped.startswith("（") and stripped.endswith("）") and len(stripped) < 80:
            add_note(doc, stripped)
            i += 1
            continue

        add_body(doc, stripped, first_line=not stripped.startswith(("**拟申报", "**所属", "**计划")))
        i += 1

    doc.save(OUT)
    print(f"wrote {OUT}  ({OUT.stat().st_size/1024:.1f} KB)")


if __name__ == "__main__":
    convert()
